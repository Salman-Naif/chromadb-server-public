"""
ChromaDB Knowledge Server
=========================
A production-ready Flask REST API for vector-based document storage and semantic search.
Perfect for RAG (Retrieval-Augmented Generation) applications.

Features:
- Document processing (PDF, DOCX, TXT)
- Vector embeddings via OpenRouter API
- Semantic similarity search
- Role-based access control
- Modern web dashboard

Author: Salman
License: MIT
Version: 1.0.0
"""

import os
import sys
import gc
import uuid
import platform
import socket
import time
import traceback
from datetime import datetime, timedelta
from functools import wraps

from flask import Flask, request, jsonify, render_template, redirect, url_for, flash, send_from_directory
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Create Flask app
app = Flask(__name__)

# ============================================
# Configuration
# ============================================

app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'change-this-secret-key-in-production')

# Database Configuration
DATABASE_URL = os.environ.get('DATABASE_URL', '')
if DATABASE_URL.startswith('postgres://'):
    DATABASE_URL = DATABASE_URL.replace('postgres://', 'postgresql://', 1)
app.config['SQLALCHEMY_DATABASE_URI'] = DATABASE_URL if DATABASE_URL else 'sqlite:///knowledge_server.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {'pool_pre_ping': True, 'pool_recycle': 300}

# Upload Configuration
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'doc', 'txt'}
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# ChromaDB Configuration
app.config['CHROMADB_PATH'] = os.environ.get('CHROMADB_PATH', './data/chromadb')
app.config['CHROMADB_COLLECTION'] = os.environ.get('CHROMADB_COLLECTION', 'knowledge_base')

# OpenRouter Configuration (for Embeddings)
app.config['OPENROUTER_API_KEY'] = os.environ.get('OPENROUTER_API_KEY', '')
app.config['EMBEDDING_MODEL'] = os.environ.get('EMBEDDING_MODEL', 'qwen/qwen3-embedding-8b')
app.config['EMBEDDING_DIMENSION'] = int(os.environ.get('EMBEDDING_DIMENSION', '4096'))

# Alternative: Google Gemini API (free tier available)
app.config['GEMINI_API_KEY'] = os.environ.get('GEMINI_API_KEY', '')

# Admin Configuration
app.config['ADMIN_USERNAME'] = os.environ.get('ADMIN_USERNAME', 'admin')
app.config['ADMIN_EMAIL'] = os.environ.get('ADMIN_EMAIL', 'admin@example.com')
app.config['ADMIN_PASSWORD'] = os.environ.get('ADMIN_PASSWORD', 'admin123')

# Flask/System info
app.config['FLASK_VERSION'] = '3.0.3'
app.config['CHROMADB_VERSION'] = '0.5.23'
app.config['ENVIRONMENT'] = os.environ.get('FLASK_ENV', 'development')

# Enable CORS
CORS(app)

# Initialize SQLAlchemy
from flask_sqlalchemy import SQLAlchemy
db = SQLAlchemy(app)

# Login Manager
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'يرجى تسجيل الدخول للوصول لهذه الصفحة'
login_manager.login_message_category = 'warning'

@login_manager.unauthorized_handler
def unauthorized():
    """Handle unauthorized access - return JSON for API, redirect for pages"""
    if request.path.startswith('/api/'):
        return jsonify({
            'success': False,
            'error': 'يرجى تسجيل الدخول',
            'redirect': url_for('login')
        }), 401
    flash('يرجى تسجيل الدخول للوصول لهذه الصفحة', 'warning')
    return redirect(url_for('login'))

# Rate Limiter
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["1000 per hour"],
    storage_uri="memory://",
)


# ============================================
# Database Models
# ============================================

class User(db.Model):
    __tablename__ = 'users'
    
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(255), unique=True, nullable=False, index=True)
    password_hash = db.Column(db.String(255), nullable=False)
    first_name = db.Column(db.String(100), default='')
    last_name = db.Column(db.String(100), default='')
    role = db.Column(db.String(50), default='staff')  # staff, admin, system_manager
    is_active = db.Column(db.Boolean, default=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    last_login = db.Column(db.DateTime)
    
    # Relationships
    activities = db.relationship('ActivityLog', backref='user', lazy='dynamic', foreign_keys='ActivityLog.user_id')
    files = db.relationship('FileMetadata', backref='uploader', lazy='dynamic', foreign_keys='FileMetadata.uploader_id')
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)
    
    @property
    def is_authenticated(self):
        return True
    
    @property
    def is_anonymous(self):
        return False
    
    def get_id(self):
        return str(self.id)
    
    @property
    def full_name(self):
        return f"{self.first_name} {self.last_name}".strip() or self.email
    
    @property
    def role_display(self):
        roles = {'staff': 'موظف', 'admin': 'مدير', 'system_manager': 'مدير النظام'}
        return roles.get(self.role, self.role)


class ActivityLog(db.Model):
    __tablename__ = 'activity_logs'
    
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=True)
    action = db.Column(db.String(100), nullable=False)
    details = db.Column(db.Text)
    ip_address = db.Column(db.String(50))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class FileMetadata(db.Model):
    __tablename__ = 'file_metadata'
    
    id = db.Column(db.Integer, primary_key=True)
    file_id = db.Column(db.String(100), unique=True, nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    file_type = db.Column(db.String(20))
    file_size = db.Column(db.Integer)
    chunks_count = db.Column(db.Integer, default=0)
    language = db.Column(db.String(10), default='ar')
    uploader_id = db.Column(db.Integer, db.ForeignKey('users.id'))
    query_count = db.Column(db.Integer, default=0)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class QueryLog(db.Model):
    __tablename__ = 'query_logs'
    
    id = db.Column(db.Integer, primary_key=True)
    query_text = db.Column(db.Text)
    results_count = db.Column(db.Integer, default=0)
    response_time_ms = db.Column(db.Float)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=True)
    source = db.Column(db.String(50), default='web')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class ErrorLog(db.Model):
    __tablename__ = 'error_logs'
    
    id = db.Column(db.Integer, primary_key=True)
    error_type = db.Column(db.String(100))
    error_message = db.Column(db.Text)
    endpoint = db.Column(db.String(255))
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


class APICallLog(db.Model):
    """Log for tracking API calls to external services"""
    __tablename__ = 'api_call_logs'
    
    id = db.Column(db.Integer, primary_key=True)
    service = db.Column(db.String(50))  # 'openrouter', 'gemini'
    endpoint = db.Column(db.String(100))
    success = db.Column(db.Boolean, default=True)
    response_time_ms = db.Column(db.Float)
    error_message = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# ============================================
# Helper Functions
# ============================================

def create_activity_log(user_id, action, details=None):
    try:
        log = ActivityLog(
            user_id=user_id,
            action=action,
            details=details,
            ip_address=request.remote_addr if request else None
        )
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        print(f"Activity log error: {e}")


def create_error_log(error_type, message, endpoint=None):
    try:
        log = ErrorLog(
            error_type=error_type,
            error_message=message,
            endpoint=endpoint or (request.path if request else None),
            user_id=current_user.id if current_user and current_user.is_authenticated else None
        )
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        print(f"Error log error: {e}")


def create_api_log(service, endpoint, success, response_time_ms, error_message=None):
    try:
        log = APICallLog(
            service=service,
            endpoint=endpoint,
            success=success,
            response_time_ms=response_time_ms,
            error_message=error_message
        )
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        print(f"API log error: {e}")


def role_required(roles):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not current_user.is_authenticated:
                flash('يرجى تسجيل الدخول', 'warning')
                return redirect(url_for('login'))
            if current_user.role not in roles:
                flash('غير مصرح لك بالوصول لهذه الصفحة', 'error')
                return redirect(url_for('dashboard'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator


def api_key_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Get API key from various sources
        api_key = None
        
        # 1. Authorization header (Bearer token)
        auth_header = request.headers.get('Authorization', '')
        if auth_header.startswith('Bearer '):
            api_key = auth_header[7:]
        
        # 2. X-API-Key header
        if not api_key:
            api_key = request.headers.get('X-API-Key', '')
        
        # 3. Query parameter
        if not api_key:
            api_key = request.args.get('api_key', '')
        
        # 4. JSON body
        if not api_key:
            data = request.get_json(silent=True) or {}
            api_key = data.get('api_key', '')
        
        # Validate
        valid_key = app.config.get('CHROMADB_API_KEY', '')
        
        if not api_key:
            return jsonify({
                'success': False, 
                'error': 'API key required',
                'hint': 'Provide via Authorization header (Bearer token), X-API-Key header, or api_key parameter'
            }), 401
        
        if api_key != valid_key:
            return jsonify({'success': False, 'error': 'Invalid API key'}), 401
        
        return f(*args, **kwargs)
    return decorated_function


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


# ============================================
# ChromaDB Manager
# ============================================

class ChromaDBManager:
    def __init__(self, path, collection_name):
        self.path = path
        self.collection_name = collection_name
        self.client = None
        self.collection = None
        self._initialize()
    
    def _initialize(self):
        try:
            import chromadb
            os.makedirs(self.path, exist_ok=True)
            self.client = chromadb.PersistentClient(path=self.path)
            self.collection = self.client.get_or_create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            print(f"✅ ChromaDB initialized: {self.collection.count()} documents")
        except Exception as e:
            print(f"❌ ChromaDB init error: {e}")
            raise
    
    def add_documents(self, chunks, embeddings, metadata_list):
        try:
            file_id = str(uuid.uuid4())
            ids = [f"{file_id}_{i}" for i in range(len(chunks))]
            
            cleaned_metadata = []
            for i, meta in enumerate(metadata_list):
                cleaned = {
                    'chunk_index': i,
                    'file_id': file_id,
                    'source': str(meta.get('source', 'unknown'))[:200],
                    'upload_date': meta.get('upload_date', datetime.now().isoformat()),
                    'lang': meta.get('lang', 'ar'),
                    'text_preview': str(meta.get('text_preview', ''))[:200]
                }
                cleaned_metadata.append(cleaned)
            
            self.collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=chunks,
                metadatas=cleaned_metadata
            )
            
            return {'success': True, 'file_id': file_id, 'chunks_count': len(chunks)}
        except Exception as e:
            print(f"❌ Add documents error: {e}")
            return {'success': False, 'error': str(e)}
    
    def search(self, query_embedding, top_k=5):
        try:
            start_time = time.time()
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=min(top_k, 10),
                include=['metadatas', 'documents', 'distances']
            )
            query_time = (time.time() - start_time) * 1000
            
            if not results.get('documents') or len(results['documents'][0]) == 0:
                return {
                    'success': True, 
                    'results': {'documents': [], 'metadatas': [], 'distances': []},
                    'query_time_ms': query_time
                }
            
            return {
                'success': True,
                'results': {
                    'documents': results['documents'][0],
                    'metadatas': results['metadatas'][0],
                    'distances': results['distances'][0]
                },
                'query_time_ms': query_time
            }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def delete_by_file_id(self, file_id):
        try:
            results = self.collection.get(where={"file_id": file_id})
            if results and results.get('ids'):
                self.collection.delete(ids=results['ids'])
                return {'success': True, 'deleted_count': len(results['ids'])}
            return {'success': False, 'error': 'File not found'}
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def get_all_files(self):
        try:
            results = self.collection.get(include=['metadatas'])
            files = {}
            for meta in results.get('metadatas', []):
                file_id = meta.get('file_id', 'unknown')
                if file_id not in files:
                    files[file_id] = {
                        'file_id': file_id,
                        'filename': meta.get('source', 'غير معروف'),
                        'lang': meta.get('lang', 'ar'),
                        'upload_date': meta.get('upload_date', ''),
                        'chunks_count': 0
                    }
                files[file_id]['chunks_count'] += 1
            
            return sorted(files.values(), key=lambda x: x.get('upload_date', ''), reverse=True)
        except Exception as e:
            print(f"Get files error: {e}")
            return []
    
    def get_file_by_id(self, file_id):
        try:
            results = self.collection.get(where={"file_id": file_id}, include=['metadatas', 'documents'])
            if not results.get('metadatas'):
                return None
            
            meta = results['metadatas'][0]
            return {
                'file_id': file_id,
                'filename': meta.get('source', ''),
                'lang': meta.get('lang', 'ar'),
                'chunks_count': len(results['metadatas']),
                'chunks': [{'text': doc[:500], 'preview': m.get('text_preview', '')} 
                          for doc, m in zip(results['documents'], results['metadatas'])]
            }
        except Exception as e:
            return None
    
    def get_stats(self):
        try:
            return {
                'total_chunks': self.collection.count(),
                'total_files': len(self.get_all_files())
            }
        except:
            return {'total_chunks': 0, 'total_files': 0}
    
    def health_check(self):
        try:
            count = self.collection.count()
            return {
                'connected': True, 
                'status': 'healthy', 
                'document_count': count,
                'collection_name': self.collection_name
            }
        except:
            return {'connected': False, 'status': 'error', 'document_count': 0, 'collection_name': ''}


# ============================================
# Embeddings Manager (OpenRouter)
# ============================================

class EmbeddingsManager:
    def __init__(self, api_key, model="qwen/qwen3-embedding-8b", dimensions=4096):
        self.api_key = api_key
        self.model = model
        self.dimensions = dimensions
        self.base_url = "https://openrouter.ai/api/v1"
        self._initialized = False
        self._initialize()
    
    def _initialize(self):
        if self.api_key:
            self._initialized = True
            print(f"✅ OpenRouter Embeddings configured: {self.model}")
        else:
            print("⚠️ OpenRouter API key not set - embeddings disabled")
    
    def generate_embedding(self, text, task_type="retrieval_document"):
        if not self._initialized or not text:
            print(f"⚠️ Embedding skipped: initialized={self._initialized}, text_len={len(text) if text else 0}")
            return None
        
        try:
            import requests
            
            # Truncate text if too long
            text = text[:8000] if len(text) > 8000 else text
            
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://github.com/chromadb-knowledge-server",
                "X-Title": "ChromaDB Knowledge Server"
            }
            
            # Try without dimensions first (some models don't support it)
            data = {
                "model": self.model,
                "input": text
            }
            
            start_time = time.time()
            response = requests.post(
                f"{self.base_url}/embeddings",
                headers=headers,
                json=data,
                timeout=60
            )
            response_time = (time.time() - start_time) * 1000
            
            if response.status_code == 200:
                result = response.json()
                if result.get('data') and len(result['data']) > 0:
                    embedding = result['data'][0]['embedding']
                    # Update dimensions based on actual response
                    if embedding:
                        self.dimensions = len(embedding)
                    create_api_log('openrouter', 'embeddings', True, response_time)
                    return embedding
            
            # Detailed error messages
            error_text = response.text[:500]
            if response.status_code == 401:
                error_msg = "❌ API Key غير صحيح أو منتهي الصلاحية - راجع OPENROUTER_API_KEY"
            elif response.status_code == 402:
                error_msg = "❌ لا يوجد رصيد كافي في حساب OpenRouter - أضف رصيد على openrouter.ai/credits"
            elif response.status_code == 404:
                error_msg = f"❌ الموديل غير موجود: {self.model}"
            elif response.status_code == 429:
                error_msg = "❌ تم تجاوز حد الطلبات - انتظر قليلاً"
            else:
                error_msg = f"Status {response.status_code}: {error_text}"
            
            create_api_log('openrouter', 'embeddings', False, response_time, error_msg)
            print(f"🔴 Embedding error: {error_msg}")
            
            # Store last error for health check
            self._last_error = error_msg
            self._last_error_code = response.status_code
            return None
            
        except requests.exceptions.Timeout:
            error_msg = "❌ انتهت مهلة الاتصال - الخادم بطيء"
            print(f"🔴 {error_msg}")
            create_api_log('openrouter', 'embeddings', False, 0, error_msg)
            return None
        except requests.exceptions.ConnectionError as e:
            error_msg = f"❌ فشل الاتصال بـ OpenRouter: {str(e)[:100]}"
            print(f"🔴 {error_msg}")
            create_api_log('openrouter', 'embeddings', False, 0, error_msg)
            return None
        except Exception as e:
            error_msg = f"❌ خطأ غير متوقع: {str(e)}"
            print(f"🔴 {error_msg}")
            create_api_log('openrouter', 'embeddings', False, 0, error_msg)
            return None
    
    def generate_query_embedding(self, query):
        return self.generate_embedding(query, task_type="retrieval_query")
    
    def generate_embeddings_batch(self, texts, on_progress=None):
        embeddings = []
        for idx, text in enumerate(texts):
            embedding = self.generate_embedding(text)
            embeddings.append(embedding)
            if on_progress:
                on_progress(idx + 1, len(texts))
            # Rate limiting
            if idx > 0 and idx % 20 == 0:
                time.sleep(0.5)
        return embeddings
    
    def health_check(self):
        if not self._initialized:
            return {
                'connected': False, 
                'status': 'not_configured', 
                'model': self.model,
                'error': 'OPENROUTER_API_KEY غير مُعرّف في البيئة'
            }
        
        try:
            start_time = time.time()
            test = self.generate_embedding("اختبار النظام")
            response_time = (time.time() - start_time) * 1000
            
            if test is not None:
                return {
                    'connected': True, 
                    'status': 'healthy', 
                    'model': self.model,
                    'dimensions': len(test) if test else self.dimensions,
                    'response_time_ms': round(response_time, 2)
                }
            
            # Return the last error if available
            error_info = getattr(self, '_last_error', 'فشل الاختبار')
            error_code = getattr(self, '_last_error_code', 0)
            
            return {
                'connected': False, 
                'status': 'error', 
                'model': self.model,
                'error': error_info,
                'error_code': error_code
            }
        except Exception as e:
            return {
                'connected': False, 
                'status': 'error', 
                'error': str(e), 
                'model': self.model
            }


# ============================================
# File Processor
# ============================================

def extract_text_from_pdf(file_stream):
    try:
        import PyPDF2
        file_stream.seek(0)
        reader = PyPDF2.PdfReader(file_stream)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        print(f"PDF error: {e}")
        return ""


def extract_text_from_docx(file_stream):
    try:
        import docx
        doc = docx.Document(file_stream)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells if cell.text])
                if row_text:
                    text.append(row_text)
        return "\n".join(text)
    except Exception as e:
        print(f"DOCX error: {e}")
        return ""


def extract_text_from_txt(file_stream):
    try:
        for encoding in ['utf-8', 'windows-1256', 'iso-8859-1']:
            try:
                file_stream.seek(0)
                return file_stream.read().decode(encoding)
            except:
                continue
        return ""
    except Exception as e:
        print(f"TXT error: {e}")
        return ""


def chunk_text(text, chunk_size=1000, overlap=100):
    """Split text into chunks - optimized for Render Free tier (512MB RAM)"""
    if not text or len(text) < 10:
        return []
    
    # Clean text
    text = ' '.join(text.split())
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if chunk and len(chunk) >= 20:
            chunks.append(chunk)
        start = end - overlap
    
    # CRITICAL: Limit chunks to prevent OUT OF MEMORY on Render Free tier
    MAX_CHUNKS = 15  # Reduced from 30 to 15 for memory safety
    if len(chunks) > MAX_CHUNKS:
        print(f"⚠️ Limiting chunks from {len(chunks)} to {MAX_CHUNKS} to prevent memory issues")
        # Take evenly distributed chunks to maintain document coverage
        step = len(chunks) // MAX_CHUNKS
        chunks = chunks[::step][:MAX_CHUNKS]
    
    return chunks


def detect_language(text):
    import re
    arabic = len(re.findall(r'[\u0600-\u06FF]', text))
    english = len(re.findall(r'[a-zA-Z]', text))
    total = arabic + english
    if total == 0:
        return 'ar'
    ratio = arabic / total
    if ratio > 0.7:
        return 'ar'
    elif ratio < 0.3:
        return 'en'
    return 'mixed'


def process_uploaded_file(file, filename):
    try:
        import io
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        file.seek(0)
        file_bytes = file.read()
        file_stream = io.BytesIO(file_bytes)
        
        if file_ext == 'pdf':
            text = extract_text_from_pdf(file_stream)
        elif file_ext in ['docx', 'doc']:
            text = extract_text_from_docx(file_stream)
        elif file_ext == 'txt':
            text = extract_text_from_txt(file_stream)
        else:
            return {'success': False, 'error': f'نوع الملف غير مدعوم: {file_ext}'}
        
        if not text or len(text.strip()) < 10:
            return {'success': False, 'error': 'لم يتم استخراج نص من الملف'}
        
        chunks = chunk_text(text)
        if not chunks:
            return {'success': False, 'error': 'فشل تقسيم النص'}
        
        return {
            'success': True,
            'text': text,
            'chunks': chunks,
            'filename': filename,
            'file_type': file_ext,
            'file_size': len(file_bytes),
            'language': detect_language(text),
            'chunks_count': len(chunks)
        }
    except Exception as e:
        print(f"Process file error: {e}")
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


# ============================================
# Initialize Services
# ============================================

chromadb_manager = None
embeddings_manager = None
start_time = datetime.now()


def init_services():
    global chromadb_manager, embeddings_manager
    try:
        chromadb_manager = ChromaDBManager(
            path=app.config['CHROMADB_PATH'],
            collection_name=app.config['CHROMADB_COLLECTION']
        )
        
        # Use OpenRouter for embeddings
        api_key = app.config.get('OPENROUTER_API_KEY') or app.config.get('GEMINI_API_KEY')
        if api_key:
            embeddings_manager = EmbeddingsManager(
                api_key=api_key,
                model=app.config['EMBEDDING_MODEL'],
                dimensions=app.config['EMBEDDING_DIMENSION']
            )
        else:
            print("⚠️ No API key configured for embeddings")
        
        print("✅ All services initialized")
    except Exception as e:
        print(f"❌ Service init error: {e}")


def get_uptime():
    return round((datetime.now() - start_time).total_seconds() / 3600, 2)


# ============================================
# Routes - Authentication
# ============================================

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        remember = request.form.get('remember') == 'on'
        
        user = User.query.filter_by(email=email).first()
        
        if user and user.check_password(password):
            if not user.is_active:
                flash('حسابك معطل. تواصل مع مدير النظام', 'error')
                return render_template('login.html')
            
            login_user(user, remember=remember)
            user.last_login = datetime.utcnow()
            db.session.commit()
            create_activity_log(user.id, 'login', 'تسجيل دخول')
            
            next_page = request.args.get('next')
            return redirect(next_page or url_for('dashboard'))
        
        flash('بيانات الدخول غير صحيحة', 'error')
    
    return render_template('login.html')


@app.route('/logout')
@login_required
def logout():
    create_activity_log(current_user.id, 'logout', 'تسجيل خروج')
    logout_user()
    flash('تم تسجيل الخروج بنجاح', 'success')
    return redirect(url_for('login'))


# ============================================
# Routes - Dashboard
# ============================================

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))


@app.route('/dashboard')
@login_required
def dashboard():
    stats = {
        'total_chunks': 0,
        'total_files': 0,
        'total_queries': QueryLog.query.count() if QueryLog.query else 0,
        'uptime': get_uptime(),
        'recent_files': [],
        'chromadb_status': False,
        'embeddings_status': False,
    }
    
    if chromadb_manager:
        try:
            chroma_stats = chromadb_manager.get_stats()
            stats['total_chunks'] = chroma_stats.get('total_chunks', 0)
            stats['total_files'] = chroma_stats.get('total_files', 0)
            stats['recent_files'] = chromadb_manager.get_all_files()[:10]
            stats['chromadb_status'] = True
        except:
            pass
    
    if embeddings_manager:
        try:
            health = embeddings_manager.health_check()
            stats['embeddings_status'] = health.get('connected', False)
        except:
            pass
    
    if current_user.role in ['admin', 'system_manager']:
        stats['total_users'] = User.query.count()
        stats['active_users'] = User.query.filter_by(is_active=True).count()
        stats['total_errors'] = ErrorLog.query.filter(
            ErrorLog.created_at >= datetime.utcnow() - timedelta(days=7)
        ).count()
    
    return render_template('dashboard.html', stats=stats)


# ============================================
# Routes - Upload
# ============================================

@app.route('/upload')
@login_required
def upload():
    return render_template('upload.html', config=app.config)


@app.route('/api/upload', methods=['POST'])
@login_required
def api_upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'لم يتم إرفاق ملف'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'success': False, 'error': 'لم يتم اختيار ملف'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'error': 'نوع الملف غير مدعوم'}), 400
        
        filename = secure_filename(file.filename)
        
        # Process file
        result = process_uploaded_file(file, filename)
        
        if not result.get('success'):
            return jsonify(result), 400
        
        # Check services
        if not chromadb_manager:
            return jsonify({'success': False, 'error': 'ChromaDB غير متوفر'}), 503
        
        if not embeddings_manager:
            return jsonify({'success': False, 'error': 'خدمة التضمينات غير متوفرة'}), 503
        
        # Generate embeddings - with memory optimization
        chunks = result['chunks']
        embeddings = []
        
        print(f"📄 Processing {len(chunks)} chunks for {filename}")
        
        # Process one at a time to save memory
        for i, chunk in enumerate(chunks):
            print(f"⏳ Embedding {i+1}/{len(chunks)}")
            
            emb = embeddings_manager.generate_embedding(chunk)
            if emb:
                embeddings.append(emb)
            else:
                embeddings.append(None)
                # If too many failures, stop early
                failed_count = sum(1 for e in embeddings if e is None)
                if failed_count > 3:
                    print(f"❌ Too many embedding failures ({failed_count}), stopping")
                    break
            
            # Force garbage collection every 5 chunks to free memory
            if i > 0 and i % 5 == 0:
                gc.collect()
        
        # Filter valid
        valid_data = [(c, e) for c, e in zip(chunks, embeddings) if e is not None]
        
        if not valid_data:
            # Get detailed error from embeddings manager
            health = embeddings_manager.health_check()
            error_detail = health.get('error', 'خطأ غير معروف')
            error_code = health.get('error_code', 0)
            
            if error_code == 401:
                error_msg = 'مفتاح API غير صحيح - راجع إعدادات OPENROUTER_API_KEY'
            elif error_code == 402:
                error_msg = 'لا يوجد رصيد كافي - أضف رصيد على openrouter.ai/credits'
            elif error_code == 404:
                error_msg = f'الموديل غير موجود: {embeddings_manager.model}'
            else:
                error_msg = f'فشل إنشاء التضمينات: {error_detail}'
            
            return jsonify({
                'success': False, 
                'error': error_msg,
                'error_code': error_code,
                'details': error_detail
            }), 500
        
        valid_chunks, valid_embeddings = zip(*valid_data)
        
        # Metadata
        metadata_list = [
            {
                'chunk_index': i,
                'source': filename,
                'upload_date': datetime.now().isoformat(),
                'lang': result.get('language', 'ar'),
                'text_preview': chunk[:200]
            }
            for i, chunk in enumerate(valid_chunks)
        ]
        
        # Store
        store_result = chromadb_manager.add_documents(
            list(valid_chunks),
            list(valid_embeddings),
            metadata_list
        )
        
        if not store_result.get('success'):
            return jsonify(store_result), 500
        
        # Save metadata
        file_meta = FileMetadata(
            file_id=store_result['file_id'],
            filename=filename,
            file_type=result.get('file_type', ''),
            file_size=result.get('file_size', 0),
            chunks_count=len(valid_chunks),
            language=result.get('language', 'ar'),
            uploader_id=current_user.id
        )
        db.session.add(file_meta)
        db.session.commit()
        
        create_activity_log(current_user.id, 'upload', f'رفع: {filename}')
        
        return jsonify({
            'success': True,
            'file_id': store_result['file_id'],
            'filename': filename,
            'chunks_count': len(valid_chunks),
            'message': f'تم رفع الملف بنجاح! {len(valid_chunks)} قطعة'
        })
        
    except Exception as e:
        print(f"Upload error: {e}")
        traceback.print_exc()
        create_error_log('UploadError', str(e))
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/upload/google-docs', methods=['POST'])
@login_required
def api_upload_google_docs():
    try:
        import io
        import re
        import requests as req
        
        data = request.get_json()
        doc_url = data.get('url', '').strip()
        
        if not doc_url:
            return jsonify({'success': False, 'error': 'لم يتم إدخال رابط'}), 400
        
        # Extract doc ID
        doc_id = None
        if 'docs.google.com' in doc_url:
            match = re.search(r'/document/d/([a-zA-Z0-9-_]+)', doc_url)
            if match:
                doc_id = match.group(1)
        elif 'drive.google.com' in doc_url:
            match = re.search(r'/file/d/([a-zA-Z0-9-_]+)', doc_url)
            if match:
                doc_id = match.group(1)
        
        if not doc_id:
            return jsonify({'success': False, 'error': 'رابط غير صحيح'}), 400
        
        # Download
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=docx"
        response = req.get(export_url, timeout=30)
        
        if response.status_code != 200:
            return jsonify({'success': False, 'error': 'فشل التحميل - تأكد من صلاحيات المشاركة'}), 400
        
        # Process
        filename = f"GoogleDoc_{doc_id[:8]}.docx"
        file_stream = io.BytesIO(response.content)
        
        class FakeFile:
            def __init__(self, stream, name):
                self.stream = stream
                self.filename = name
            def read(self):
                return self.stream.read()
            def seek(self, pos):
                self.stream.seek(pos)
        
        fake_file = FakeFile(file_stream, filename)
        result = process_uploaded_file(fake_file, filename)
        
        if not result.get('success'):
            return jsonify(result), 400
        
        # Same as regular upload
        if not embeddings_manager:
            return jsonify({'success': False, 'error': 'خدمة التضمينات غير متوفرة'}), 503
        
        chunks = result['chunks']
        embeddings = [embeddings_manager.generate_embedding(c) for c in chunks]
        valid_data = [(c, e) for c, e in zip(chunks, embeddings) if e]
        
        if not valid_data:
            return jsonify({'success': False, 'error': 'فشل إنشاء التضمينات'}), 500
        
        valid_chunks, valid_embeddings = zip(*valid_data)
        
        metadata_list = [
            {'source': filename, 'upload_date': datetime.now().isoformat(), 'lang': result.get('language', 'ar'), 'text_preview': c[:200]}
            for c in valid_chunks
        ]
        
        store_result = chromadb_manager.add_documents(list(valid_chunks), list(valid_embeddings), metadata_list)
        
        if store_result.get('success'):
            create_activity_log(current_user.id, 'upload', f'Google Doc: {filename}')
            
            # Save metadata
            file_meta = FileMetadata(
                file_id=store_result['file_id'],
                filename=filename,
                file_type='docx',
                file_size=len(response.content),
                chunks_count=len(valid_chunks),
                language=result.get('language', 'ar'),
                uploader_id=current_user.id
            )
            db.session.add(file_meta)
            db.session.commit()
        
        return jsonify(store_result)
        
    except Exception as e:
        print(f"Google Docs error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================
# Routes - File Manager
# ============================================

@app.route('/files')
@login_required
def file_manager():
    files = chromadb_manager.get_all_files() if chromadb_manager else []
    return render_template('file_manager.html', files=files)


@app.route('/api/file/<file_id>', methods=['GET'])
@login_required
def api_get_file(file_id):
    try:
        file_data = chromadb_manager.get_file_by_id(file_id) if chromadb_manager else None
        if file_data:
            return jsonify({'success': True, 'file': file_data})
        return jsonify({'success': False, 'error': 'الملف غير موجود'}), 404
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/file/<file_id>', methods=['DELETE'])
@login_required
@role_required(['admin', 'system_manager'])
def api_delete_file(file_id):
    try:
        result = chromadb_manager.delete_by_file_id(file_id) if chromadb_manager else {'success': False}
        if result.get('success'):
            FileMetadata.query.filter_by(file_id=file_id).delete()
            db.session.commit()
            create_activity_log(current_user.id, 'delete', f'حذف: {file_id}')
        return jsonify(result)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================
# Routes - Search API
# ============================================

@app.route('/search', methods=['POST'])
@api_key_required
def api_search():
    try:
        data = request.get_json()
        query_text = data.get('query', '').strip()
        top_k = min(data.get('topK', 5), 10)
        
        if not query_text:
            return jsonify({'success': False, 'error': 'لم يتم إدخال استعلام'}), 400
        
        if not embeddings_manager or not chromadb_manager:
            return jsonify({'success': False, 'error': 'الخدمات غير متوفرة'}), 503
        
        query_embedding = embeddings_manager.generate_query_embedding(query_text)
        
        if not query_embedding:
            return jsonify({'success': False, 'error': 'فشل إنشاء التضمين - تحقق من مفتاح API'}), 500
        
        start = datetime.now()
        results = chromadb_manager.search(query_embedding, top_k)
        response_time = (datetime.now() - start).total_seconds() * 1000
        
        # Log
        query_log = QueryLog(
            query_text=query_text[:500],
            results_count=len(results.get('results', {}).get('documents', [])),
            response_time_ms=response_time,
            source='api'
        )
        db.session.add(query_log)
        db.session.commit()
        
        return jsonify(results)
        
    except Exception as e:
        create_error_log('SearchError', str(e))
        return jsonify({'success': False, 'error': str(e)}), 500


# ============================================
# Routes - Status & Health
# ============================================

@app.route('/status')
@login_required
@role_required(['admin', 'system_manager'])
def status():
    import flask
    
    chromadb_health = chromadb_manager.health_check() if chromadb_manager else {'connected': False, 'status': 'غير متصل', 'document_count': 0, 'collection_name': ''}
    embeddings_health = embeddings_manager.health_check() if embeddings_manager else {'connected': False, 'status': 'غير مهيأ'}
    
    status_data = {
        'chromadb': {
            'connected': chromadb_health.get('connected', False),
            'status': chromadb_health.get('status', 'غير متصل'),
            'collection_name': app.config.get('CHROMADB_COLLECTION', 'غير محدد'),
            'document_count': chromadb_health.get('document_count', 0)
        },
        'embeddings': embeddings_health,
        'system': {
            'python_version': platform.python_version(),
            'flask_version': flask.__version__,
            'chromadb_version': app.config.get('CHROMADB_VERSION', '0.5.23'),
            'environment': app.config.get('ENVIRONMENT', 'development'),
            'hostname': socket.gethostname(),
            'uptime': get_uptime(),
            'start_time': start_time.strftime('%Y-%m-%d %H:%M')
        }
    }
    
    activities = ActivityLog.query.order_by(ActivityLog.created_at.desc()).limit(50).all()
    return render_template('status.html', status=status_data, activities=activities)


@app.route('/health')
@login_required
@role_required(['admin', 'system_manager'])
def health():
    try:
        import psutil
        cpu = psutil.cpu_percent()
        ram = psutil.virtual_memory().percent
        disk = psutil.disk_usage('/').percent
    except:
        cpu = ram = disk = 0
    
    # Measure database latency
    latency = 0
    try:
        start = time.time()
        User.query.first()
        latency = int((time.time() - start) * 1000)
    except:
        pass
    
    # Database stats
    failed_queries = ErrorLog.query.filter(
        ErrorLog.created_at >= datetime.utcnow() - timedelta(days=1),
        ErrorLog.error_type.like('%Search%')
    ).count()
    
    avg_query_time = 0
    try:
        result = db.session.query(db.func.avg(QueryLog.response_time_ms)).filter(
            QueryLog.created_at >= datetime.utcnow() - timedelta(days=1)
        ).scalar()
        avg_query_time = round(result or 0, 2)
    except:
        pass
    
    # API stats
    api_failed_calls = APICallLog.query.filter(
        APICallLog.created_at >= datetime.utcnow() - timedelta(days=1),
        APICallLog.success == False
    ).count()
    
    api_response_time = 0
    try:
        result = db.session.query(db.func.avg(APICallLog.response_time_ms)).filter(
            APICallLog.created_at >= datetime.utcnow() - timedelta(days=1),
            APICallLog.success == True
        ).scalar()
        api_response_time = round(result or 0, 2)
    except:
        pass
    
    # Embeddings status
    embeddings_health = embeddings_manager.health_check() if embeddings_manager else {'connected': False, 'status': 'غير مهيأ'}
    
    health_data = {
        'system': {
            'cpu': cpu, 
            'ram': ram, 
            'disk': disk, 
            'latency': latency
        },
        'database': {
            'collections': 1,
            'avg_query_time': avg_query_time,
            'failed_queries': failed_queries,
            'index_status': 'سليم'
        },
        'api': {
            'status': embeddings_health.get('status', 'غير معروف'),
            'response_time': api_response_time,
            'failed_calls': api_failed_calls,
            'rate_limit_status': 'طبيعي',
            'model': embeddings_health.get('model', 'غير محدد'),
            'error': embeddings_health.get('error'),
            'error_code': embeddings_health.get('error_code', 0)
        },
        'alerts': []
    }
    
    # Generate alerts
    if cpu > 80:
        health_data['alerts'].append({'type': 'warning', 'title': 'CPU عالي', 'message': f'استخدام المعالج {cpu}%'})
    if ram > 85:
        health_data['alerts'].append({'type': 'warning', 'title': 'ذاكرة عالية', 'message': f'استخدام الذاكرة {ram}%'})
    if disk > 85:
        health_data['alerts'].append({'type': 'error', 'title': 'تخزين منخفض', 'message': f'المساحة المستخدمة {disk}%'})
    if failed_queries > 10:
        health_data['alerts'].append({'type': 'warning', 'title': 'أخطاء بحث', 'message': f'{failed_queries} استعلام فاشل في 24 ساعة'})
    if not embeddings_health.get('connected'):
        health_data['alerts'].append({'type': 'error', 'title': 'API غير متصل', 'message': 'خدمة التضمينات غير متاحة'})
    
    errors = ErrorLog.query.filter(
        ErrorLog.created_at >= datetime.utcnow() - timedelta(days=1)
    ).order_by(ErrorLog.created_at.desc()).limit(20).all()
    
    return render_template('health.html', health=health_data, errors=errors)


# ============================================
# Routes - Profile
# ============================================

@app.route('/profile')
@login_required
def profile():
    activities = ActivityLog.query.filter_by(user_id=current_user.id).order_by(ActivityLog.created_at.desc()).limit(10).all()
    return render_template('profile.html', activities=activities)


@app.route('/profile/update', methods=['POST'])
@login_required
def update_profile():
    current_user.first_name = request.form.get('first_name', '').strip()
    current_user.last_name = request.form.get('last_name', '').strip()
    db.session.commit()
    flash('تم تحديث المعلومات', 'success')
    return redirect(url_for('profile'))


@app.route('/profile/change-password', methods=['POST'])
@login_required
def change_password():
    current_password = request.form.get('current_password')
    new_password = request.form.get('new_password')
    confirm_password = request.form.get('confirm_password')
    
    if not current_user.check_password(current_password):
        flash('كلمة المرور الحالية غير صحيحة', 'error')
        return redirect(url_for('profile'))
    
    if new_password != confirm_password:
        flash('كلمة المرور غير متطابقة', 'error')
        return redirect(url_for('profile'))
    
    if len(new_password) < 8:
        flash('كلمة المرور قصيرة جداً (8 أحرف على الأقل)', 'error')
        return redirect(url_for('profile'))
    
    current_user.set_password(new_password)
    db.session.commit()
    flash('تم تغيير كلمة المرور', 'success')
    return redirect(url_for('profile'))


# ============================================
# Routes - User Management
# ============================================

@app.route('/users')
@login_required
@role_required(['admin', 'system_manager'])
def user_management():
    users = User.query.order_by(User.created_at.desc()).all()
    stats = {
        'total': len(users),
        'active': len([u for u in users if u.is_active]),
        'admins': len([u for u in users if u.role == 'admin']),
        'managers': len([u for u in users if u.role == 'system_manager'])
    }
    return render_template('users.html', users=users, stats=stats)


@app.route('/users/add', methods=['POST'])
@login_required
@role_required(['system_manager'])
def add_user():
    email = request.form.get('email', '').strip().lower()
    first_name = request.form.get('first_name', '').strip()
    last_name = request.form.get('last_name', '').strip()
    password = request.form.get('password')
    role = request.form.get('role', 'staff')
    
    if not all([email, password]):
        flash('البريد وكلمة المرور مطلوبان', 'error')
        return redirect(url_for('user_management'))
    
    if len(password) < 8:
        flash('كلمة المرور يجب أن تكون 8 أحرف على الأقل', 'error')
        return redirect(url_for('user_management'))
    
    if User.query.filter_by(email=email).first():
        flash('البريد مسجل مسبقاً', 'error')
        return redirect(url_for('user_management'))
    
    user = User(email=email, first_name=first_name, last_name=last_name, role=role)
    user.set_password(password)
    db.session.add(user)
    db.session.commit()
    
    create_activity_log(current_user.id, 'add_user', f'إضافة مستخدم: {email}')
    flash(f'تم إضافة {email}', 'success')
    return redirect(url_for('user_management'))


@app.route('/users/update', methods=['POST'])
@login_required
@role_required(['system_manager'])
def update_user():
    user_id = request.form.get('user_id')
    user = User.query.get(user_id)
    
    if not user:
        flash('المستخدم غير موجود', 'error')
        return redirect(url_for('user_management'))
    
    user.first_name = request.form.get('first_name', '').strip()
    user.last_name = request.form.get('last_name', '').strip()
    user.role = request.form.get('role', user.role)
    
    new_password = request.form.get('new_password')
    if new_password and len(new_password) >= 8:
        user.set_password(new_password)
    
    db.session.commit()
    create_activity_log(current_user.id, 'update_user', f'تعديل مستخدم: {user.email}')
    flash('تم التحديث', 'success')
    return redirect(url_for('user_management'))


@app.route('/api/user/<int:user_id>/toggle-status', methods=['POST'])
@login_required
@role_required(['system_manager'])
def toggle_user_status(user_id):
    user = User.query.get(user_id)
    if not user:
        return jsonify({'success': False, 'error': 'المستخدم غير موجود'}), 404
    if user.id == current_user.id:
        return jsonify({'success': False, 'error': 'لا يمكنك تعطيل نفسك'}), 400
    
    user.is_active = not user.is_active
    db.session.commit()
    
    action = 'تفعيل' if user.is_active else 'تعطيل'
    create_activity_log(current_user.id, 'toggle_user', f'{action} مستخدم: {user.email}')
    
    return jsonify({'success': True, 'is_active': user.is_active})


@app.route('/api/user/<int:user_id>/delete', methods=['DELETE'])
@login_required
@role_required(['system_manager'])
def delete_user(user_id):
    user = User.query.get(user_id)
    if not user:
        return jsonify({'success': False, 'error': 'المستخدم غير موجود'}), 404
    if user.id == current_user.id:
        return jsonify({'success': False, 'error': 'لا يمكنك حذف نفسك'}), 400
    if user.role == 'system_manager':
        # Check if this is the last system_manager
        managers_count = User.query.filter_by(role='system_manager').count()
        if managers_count <= 1:
            return jsonify({'success': False, 'error': 'لا يمكن حذف آخر مدير نظام'}), 400
    
    email = user.email
    db.session.delete(user)
    db.session.commit()
    
    create_activity_log(current_user.id, 'delete_user', f'حذف مستخدم: {email}')
    
    return jsonify({'success': True, 'message': f'تم حذف {email}'})


# ============================================
# Routes - Analysis
# ============================================

@app.route('/analysis')
@login_required
@role_required(['admin', 'system_manager'])
def analysis():
    # Get query stats
    total_queries = QueryLog.query.count()
    
    # Get queries per day for the last 7 days
    seven_days_ago = datetime.utcnow() - timedelta(days=7)
    daily_queries = db.session.query(
        db.func.date(QueryLog.created_at).label('date'),
        db.func.count(QueryLog.id).label('count')
    ).filter(QueryLog.created_at >= seven_days_ago).group_by(
        db.func.date(QueryLog.created_at)
    ).all()
    
    chart_data = [{'date': str(d.date), 'count': d.count} for d in daily_queries]
    
    # Get queries by hour
    hourly_queries = [0] * 24
    try:
        results = db.session.query(
            db.func.extract('hour', QueryLog.created_at).label('hour'),
            db.func.count(QueryLog.id).label('count')
        ).filter(QueryLog.created_at >= seven_days_ago).group_by(
            db.func.extract('hour', QueryLog.created_at)
        ).all()
        for r in results:
            if r.hour is not None:
                hourly_queries[int(r.hour)] = r.count
    except:
        pass
    
    analytics = {
        'queries': {
            'total': total_queries,
            'avg_per_day': round(total_queries / max(get_uptime() / 24, 1), 1),
            'success_rate': 95,
            'chart_data': chart_data,
            'hours': hourly_queries
        },
        'documents': {
            'most_accessed': [],
            'types': {'PDF': 0, 'DOCX': 0, 'TXT': 0},
            'language_dist': {'ar': 0, 'en': 0, 'mixed': 0, 'ar_percent': 0, 'en_percent': 0, 'mixed_percent': 0}
        },
        'performance': {'avg_retrieval_time': 50},
        'users': []
    }
    
    if chromadb_manager:
        files = chromadb_manager.get_all_files()
        for f in files:
            ext = f['filename'].rsplit('.', 1)[-1].upper() if '.' in f['filename'] else 'TXT'
            if ext in analytics['documents']['types']:
                analytics['documents']['types'][ext] += 1
            
            lang = f.get('lang', 'ar')
            if lang in analytics['documents']['language_dist']:
                analytics['documents']['language_dist'][lang] += 1
        
        total_files = sum(analytics['documents']['types'].values())
        if total_files > 0:
            for lang in ['ar', 'en', 'mixed']:
                analytics['documents']['language_dist'][f'{lang}_percent'] = round(
                    analytics['documents']['language_dist'][lang] / total_files * 100, 1
                )
    
    return render_template('analysis.html', analytics=analytics)


# ============================================
# API Endpoints
# ============================================

@app.route('/api/health-check')
def api_health_check():
    chromadb_ok = chromadb_manager.health_check().get('connected', False) if chromadb_manager else False
    embeddings_ok = embeddings_manager.health_check().get('connected', False) if embeddings_manager else False
    
    return jsonify({
        'status': 'healthy' if chromadb_ok else 'degraded',
        'timestamp': datetime.now().isoformat(),
        'version': '3.0.0',
        'services': {
            'chromadb': chromadb_ok,
            'embeddings': embeddings_ok
        }
    })


@app.route('/wake-up')
def wake_up():
    doc_count = chromadb_manager.get_stats().get('total_chunks', 0) if chromadb_manager else 0
    return jsonify({
        'status': 'awake', 
        'documents': doc_count, 
        'timestamp': datetime.now().isoformat()
    })


@app.route('/api/stats')
@api_key_required
def api_stats():
    """Get system statistics via API"""
    stats = {
        'total_chunks': 0,
        'total_files': 0,
        'total_queries': QueryLog.query.count(),
        'uptime_hours': get_uptime()
    }
    
    if chromadb_manager:
        chroma_stats = chromadb_manager.get_stats()
        stats['total_chunks'] = chroma_stats.get('total_chunks', 0)
        stats['total_files'] = chroma_stats.get('total_files', 0)
    
    return jsonify({'success': True, 'stats': stats})


@app.route('/api/diagnose')
@login_required
@role_required(['admin', 'system_manager'])
def api_diagnose():
    """Comprehensive system diagnosis endpoint"""
    import requests
    
    diagnosis = {
        'timestamp': datetime.now().isoformat(),
        'environment': {},
        'services': {},
        'tests': {}
    }
    
    # Environment check
    diagnosis['environment'] = {
        'OPENROUTER_API_KEY': '✅ موجود' if app.config.get('OPENROUTER_API_KEY') else '❌ غير موجود',
        'OPENROUTER_API_KEY_preview': app.config.get('OPENROUTER_API_KEY', '')[:20] + '...' if app.config.get('OPENROUTER_API_KEY') else 'N/A',
        'EMBEDDING_MODEL': app.config.get('EMBEDDING_MODEL', 'غير محدد'),
        'EMBEDDING_DIMENSION': app.config.get('EMBEDDING_DIMENSION', 0),
        'DATABASE_URL': '✅ موجود' if app.config.get('SQLALCHEMY_DATABASE_URI') else '❌ غير موجود'
    }
    
    # ChromaDB check
    if chromadb_manager:
        diagnosis['services']['chromadb'] = chromadb_manager.health_check()
    else:
        diagnosis['services']['chromadb'] = {'status': 'not_initialized'}
    
    # Embeddings check
    if embeddings_manager:
        diagnosis['services']['embeddings'] = embeddings_manager.health_check()
    else:
        diagnosis['services']['embeddings'] = {'status': 'not_initialized'}
    
    # Direct API test
    api_key = app.config.get('OPENROUTER_API_KEY')
    if api_key:
        try:
            headers = {
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            }
            
            # Test with simple request
            test_response = requests.post(
                "https://openrouter.ai/api/v1/embeddings",
                headers=headers,
                json={
                    "model": app.config.get('EMBEDDING_MODEL', 'qwen/qwen3-embedding-8b'),
                    "input": "test"
                },
                timeout=30
            )
            
            diagnosis['tests']['direct_api'] = {
                'status_code': test_response.status_code,
                'success': test_response.status_code == 200,
                'response_preview': test_response.text[:300] if test_response.status_code != 200 else 'OK'
            }
            
            # Interpret error
            if test_response.status_code == 401:
                diagnosis['tests']['direct_api']['diagnosis'] = 'مفتاح API غير صحيح أو منتهي'
                diagnosis['tests']['direct_api']['solution'] = 'أنشئ مفتاح جديد من openrouter.ai/keys'
            elif test_response.status_code == 402:
                diagnosis['tests']['direct_api']['diagnosis'] = 'لا يوجد رصيد في الحساب'
                diagnosis['tests']['direct_api']['solution'] = 'أضف رصيد من openrouter.ai/credits'
            elif test_response.status_code == 404:
                diagnosis['tests']['direct_api']['diagnosis'] = 'الموديل غير موجود'
                diagnosis['tests']['direct_api']['solution'] = 'تحقق من اسم الموديل'
                
        except requests.exceptions.Timeout:
            diagnosis['tests']['direct_api'] = {
                'success': False,
                'error': 'Timeout - الخادم بطيء'
            }
        except Exception as e:
            diagnosis['tests']['direct_api'] = {
                'success': False,
                'error': str(e)
            }
    else:
        diagnosis['tests']['direct_api'] = {
            'success': False,
            'error': 'No API key configured'
        }
    
    # Recent API errors
    recent_errors = APICallLog.query.filter(
        APICallLog.success == False,
        APICallLog.created_at >= datetime.utcnow() - timedelta(hours=24)
    ).order_by(APICallLog.created_at.desc()).limit(10).all()
    
    diagnosis['recent_errors'] = [
        {
            'time': e.created_at.strftime('%H:%M:%S'),
            'service': e.service,
            'error': e.error_message[:100] if e.error_message else 'Unknown'
        }
        for e in recent_errors
    ]
    
    return jsonify(diagnosis)


# ============================================
# Error Handlers
# ============================================

@app.errorhandler(404)
def not_found(e):
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'error': 'Endpoint not found'}), 404
    flash('الصفحة غير موجودة', 'error')
    return redirect(url_for('dashboard'))


@app.errorhandler(500)
def server_error(e):
    create_error_log('ServerError', str(e))
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'error': 'Internal server error'}), 500
    flash('خطأ في الخادم', 'error')
    return redirect(url_for('dashboard'))


@app.errorhandler(413)
def too_large(e):
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'error': 'File too large (max 50MB)'}), 413
    flash('حجم الملف كبير جداً (الحد الأقصى 50 ميجابايت)', 'error')
    return redirect(url_for('upload'))


# ============================================
# Application Initialization
# ============================================

def init_app():
    with app.app_context():
        db.create_all()
        
        admin_email = app.config['ADMIN_EMAIL']
        admin = User.query.filter_by(email=admin_email).first()
        
        if not admin:
            admin = User(
                email=admin_email,
                first_name='مدير',
                last_name='النظام',
                role='system_manager'
            )
            admin.set_password(app.config['ADMIN_PASSWORD'])
            db.session.add(admin)
            db.session.commit()
            print(f"✅ Admin created: {admin_email}")
        
        init_services()


# Initialize
init_app()


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8000))
    print(f"🚀 Server: http://localhost:{port}")
    app.run(host='0.0.0.0', port=port, debug=False)
